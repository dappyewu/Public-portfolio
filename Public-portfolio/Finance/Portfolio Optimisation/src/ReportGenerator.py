import re
from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Inches


# -----------------------------
# Helpers
# -----------------------------
def pick_latest_file(folder: Path, prefix: str, ext: str):
    """
    Pick latest file by timestamp pattern prefixYYYYMMDD_HHMMSS.ext
    Example: Methods_summary_20260124_150916.csv
    """
    pattern = re.compile(rf"^{re.escape(prefix)}(\d{{8}}_\d{{6}})\.{re.escape(ext)}$")
    candidates = []
    for f in folder.iterdir():
        m = pattern.match(f.name)
        if m:
            candidates.append((m.group(1), f))

    if not candidates:
        return None

    candidates.sort(key=lambda x: x[0])  # timestamp string sorts correctly
    return candidates[-1][1]


def safe_float(x, default=None):
    try:
        return float(x)
    except Exception:
        return default


def add_df_table(doc: Document, df: pd.DataFrame, title: str, max_rows=30, float_fmt="{:.4f}"):
    doc.add_heading(title, level=2)

    df_show = df.copy()
    if len(df_show) > max_rows:
        df_show = df_show.head(max_rows)

    # ensure index is a column (useful if df has meaningful index)
    df_show = df_show.reset_index(drop=False)

    table = doc.add_table(rows=1, cols=len(df_show.columns))
    table.style = "Light List Accent 1"

    # headers
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df_show.columns):
        hdr_cells[j].text = str(col)

    # rows
    for _, row in df_show.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row.values):
            if isinstance(val, (int, float)) and pd.notna(val):
                row_cells[j].text = float_fmt.format(val)
            else:
                row_cells[j].text = str(val)

    doc.add_paragraph("")


def top_weights_table(weights_df: pd.DataFrame, top_n=5):
    """
    weights_df structure like:
      Method | Ticker1 | Ticker2 | ...
    Returns compact table with top N tickers per method.
    """
    if weights_df is None or weights_df.empty:
        return None

    if "Method" not in weights_df.columns:
        return None

    out = []
    tickers = [c for c in weights_df.columns if c != "Method"]

    for _, r in weights_df.iterrows():
        method = r["Method"]
        pairs = [(t, safe_float(r[t], 0.0)) for t in tickers]
        pairs = sorted(pairs, key=lambda x: x[1], reverse=True)
        top = pairs[:top_n]

        row = {"Method": method}
        for i, (t, w) in enumerate(top, start=1):
            row[f"Top{i}"] = f"{t}: {w:.2f}%"
        out.append(row)

    return pd.DataFrame(out)


def parse_drawdown_info(text: str):
    """
    Try to extract drawdown date and % from drawdown_info file.
    Handles lots of possible formats.
    Returns: (date_str, dd_pct_float)
    """
    if not text:
        return None, None

    # find percent like -25.43% or 25.43%
    pct_match = re.search(r"(-?\d+(\.\d+)?)\s*%", text)
    dd_pct = float(pct_match.group(1)) if pct_match else None

    # find a date-like string
    date_match = re.search(r"(\d{4}-\d{2}-\d{2})", text)  # 2020-03-20
    date_str = date_match.group(1) if date_match else None

    # fallback to any "YYYY/MM/DD"
    if date_str is None:
        date_match2 = re.search(r"(\d{4}/\d{2}/\d{2})", text)
        date_str = date_match2.group(1) if date_match2 else None

    return date_str, dd_pct


# -----------------------------
# Main Word report builder
# -----------------------------
def build_word_summary(processed_folder="data/processed", output_name=None):
    folder = Path(processed_folder)
    folder.mkdir(parents=True, exist_ok=True)

    # ✅ auto-name report with timestamp + always save inside processed_folder
    if output_name is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"Backtest_Summary_Report_{ts}.docx"
    output_path = folder / output_name

    # Find latest files
    portfolio_results = pick_latest_file(folder, "portfolio_results_", "csv")
    methods_summary = pick_latest_file(folder, "Allocation_Style_summary_", "csv")
    cagr_file = pick_latest_file(folder, "CAGR_", "csv")
    drawdown_table_file = pick_latest_file(folder, "drawdown_table_", "csv")
    drawdown_info_file = pick_latest_file(folder, "drawdown_info_", "txt")

    # Pick latest idealweights (you may have multiple)
    idealweights_file = None
    ideal_candidates = sorted(folder.glob("Allocationwe*.csv"))
    if ideal_candidates:
        idealweights_file = ideal_candidates[-1]

    # Charts
    efficient_frontier_img = pick_latest_file(folder, "efficient_frontier_", "png")
    wealth_index_img = pick_latest_file(folder, "WealthIndexComparison_", "png")

    # Read data safely
    pr = pd.read_csv(portfolio_results) if portfolio_results else None
    ms = pd.read_csv(methods_summary) if methods_summary else None
    cagr = pd.read_csv(cagr_file) if cagr_file else None
    dd_table = pd.read_csv(drawdown_table_file) if drawdown_table_file else None
    iw = pd.read_csv(idealweights_file) if idealweights_file else None

    dd_info_text = ""
    if drawdown_info_file and drawdown_info_file.exists():
        dd_info_text = drawdown_info_file.read_text().strip()

    # Tidy CAGR column name
    if cagr is not None and "0" in cagr.columns:
        cagr = cagr.rename(columns={"0": "CAGR"})

    # Identify best & worst method (by Sharpe)
    best_method_line = ""
    worst_method_line = ""

    if ms is not None and "Sharpe" in ms.columns:
        ms_clean = ms.dropna(subset=["Sharpe"]).copy()
        if not ms_clean.empty:
            best_row = ms_clean.loc[ms_clean["Sharpe"].idxmax()]
            worst_row = ms_clean.loc[ms_clean["Sharpe"].idxmin()]

            best_method_line = (
                f"Best method (by Sharpe): {best_row['Method']} "
                f"(Sharpe={best_row['Sharpe']:.3f}, Return={best_row['Return (ann.%)']:.2f}%, "
                f"Vol={best_row['Vol (ann.%)']:.2f}%)"
            )
            worst_method_line = (
                f"Worst method (by Sharpe): {worst_row['Method']} "
                f"(Sharpe={worst_row['Sharpe']:.3f}, Return={worst_row['Return (ann.%)']:.2f}%, "
                f"Vol={worst_row['Vol (ann.%)']:.2f}%)"
            )

    # Top 3 tickers by CAGR
    top3_cagr_line = ""
    if cagr is not None and "Ticker" in cagr.columns and "CAGR" in cagr.columns:
        top3 = (
            cagr.dropna(subset=["CAGR"])
            .sort_values("CAGR", ascending=False)
            .head(3)
        )
        if not top3.empty:
            items = [f"{r['Ticker']} ({r['CAGR']:.2f})" for _, r in top3.iterrows()]
            top3_cagr_line = "Top 3 tickers by CAGR: " + ", ".join(items)

    # Parse drawdown info into structured
    dd_date, dd_pct = parse_drawdown_info(dd_info_text)
    dd_summary_line = ""
    if dd_date and dd_pct is not None:
        dd_summary_line = f"Worst drawdown: {dd_pct:.2f}% on {dd_date}"
    elif dd_info_text:
        dd_summary_line = f"Worst drawdown: {dd_info_text}"

    # Create document
    doc = Document()
    doc.add_heading("Backtest Summary Report", level=1)

    # Executive summary
    doc.add_heading("Executive Summary", level=2)

    # Portfolio summary line
    if pr is not None and not pr.empty:
        # your file has these columns (from your earlier bundle):
        # Portfolio Returns%, Portfolio Volatility%, Portfolio Sharpe Ratio
        ret = safe_float(pr.loc[0].get("Portfolio Returns%"))
        vol = safe_float(pr.loc[0].get("Portfolio Volatility%"))
        shr = safe_float(pr.loc[0].get("Portfolio Sharpe Ratio"))

        if ret is not None and vol is not None and shr is not None:
            doc.add_paragraph(
                f"Portfolio (annualised): Return={ret:.2f}%, Volatility={vol:.2f}%, Sharpe={shr:.3f}"
            )

    # Key takeaways bullets
    doc.add_paragraph("Key takeaways:")
    bullets = []

    if best_method_line:
        bullets.append(best_method_line)
    if worst_method_line:
        bullets.append(worst_method_line)
    if top3_cagr_line:
        bullets.append(top3_cagr_line)
    if dd_summary_line:
        bullets.append(dd_summary_line)

    if not bullets:
        bullets.append("No summary metrics were available (missing inputs in processed folder).")

    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph("")

    # Charts
    doc.add_heading("Charts", level=2)

    if wealth_index_img and wealth_index_img.exists():
        doc.add_paragraph("Wealth Index Comparison")
        doc.add_picture(str(wealth_index_img), width=Inches(6.5))

    if efficient_frontier_img and efficient_frontier_img.exists():
        doc.add_paragraph("Efficient Frontier")
        doc.add_picture(str(efficient_frontier_img), width=Inches(6.5))

    doc.add_paragraph("")

    # Tables
    if pr is not None:
        add_df_table(doc, pr, "Portfolio Results", max_rows=10, float_fmt="{:.4f}")

    if ms is not None:
        add_df_table(doc, ms, "Methods Summary", max_rows=20, float_fmt="{:.4f}")

    if iw is not None:
        compact_weights = top_weights_table(iw, top_n=5)
        if compact_weights is not None:
            add_df_table(doc, compact_weights, "Top Weights (Top 5 per Method)", max_rows=50, float_fmt="{}")

    if cagr is not None and "CAGR" in cagr.columns:
        cagr_sorted = cagr.sort_values("CAGR", ascending=False)
        add_df_table(doc, cagr_sorted.head(15), "Top 15 CAGR Tickers", max_rows=15, float_fmt="{:.4f}")

    if dd_table is not None:
        add_df_table(doc, dd_table.tail(15), "Drawdown Table (Last 15 rows)", max_rows=15, float_fmt="{:.4f}")

    # IMPORTANT: ignore WealthIndexComparisonData_*.csv (not referenced anywhere)

    doc.save(output_path)
    print(f"✅ Word report created: {output_path}")
    return output_path


# -----------------------------
# Script runner
# -----------------------------
if __name__ == "__main__":
    # ✅ This will save the report INTO data/processed automatically
    build_word_summary(processed_folder="data/processed")

import argparse
from pathlib import Path
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

# pyfolio (use pyfolio-reloaded if on pandas>=2)
import pyfolio as pf

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


APPROX_TRADING_DAYS = 252


def load_returns(path: Path) -> pd.Series:
    df = pd.read_csv(path, parse_dates=["date"]).sort_values("date").set_index("date")
    rets = df["return"].astype(float).dropna()
    rets.index = pd.to_datetime(rets.index)
    rets = rets.sort_index()
    return rets


def load_benchmark_prices(path: Path) -> pd.Series:
    df = pd.read_csv(path, parse_dates=["date"]).sort_values("date").set_index("date")
    prices = df["price"].astype(float).dropna()
    prices.index = pd.to_datetime(prices.index)
    prices = prices.sort_index()
    return prices.pct_change().dropna()


def load_risk_free_daily(path: Path) -> pd.Series:
    """Load a *daily* risk-free return series.

    Expected CSV format: `date, rf`
      - `date`: YYYY-MM-DD
      - `rf`: daily risk-free *return* as a decimal (e.g., 0.00015)

    If you only have an *annualized yield* (e.g., 5%), convert it to a daily
    return first (roughly `0.05/252`) and write that to the CSV.
    """
    df = pd.read_csv(path, parse_dates=["date"]).sort_values("date").set_index("date")
    if "rf" not in df.columns:
        raise ValueError(f"Risk-free CSV must contain an 'rf' column. Found columns: {list(df.columns)}")
    rf = df["rf"].astype(float).dropna()
    rf.index = pd.to_datetime(rf.index)
    rf = rf.sort_index()
    return rf


def calc_equity_curve(returns: pd.Series, start: float = 1.0) -> pd.Series:
    return start * (1.0 + returns).cumprod()


def calc_drawdown(returns: pd.Series) -> pd.Series:
    wealth = calc_equity_curve(returns, start=1.0)
    peak = wealth.cummax()
    dd = wealth / peak - 1.0
    return dd


def rolling_sharpe(returns: pd.Series, window: int = 252) -> pd.Series:
    # NOTE: assumes risk-free rate = 0 unless caller supplies excess returns.
    mu = returns.rolling(window, min_periods=window).mean()
    sig = returns.rolling(window, min_periods=window).std()
    sr = (mu / sig) * np.sqrt(APPROX_TRADING_DAYS)
    sr = sr.replace([np.inf, -np.inf], np.nan)
    return sr


def rolling_vol(returns: pd.Series, window: int = 63) -> pd.Series:
    return returns.rolling(window, min_periods=window).std() * np.sqrt(APPROX_TRADING_DAYS)


def rolling_beta(returns: pd.Series, benchmark_rets: pd.Series, window: int = 252) -> pd.Series:
    # beta = cov(r, b) / var(b)
    cov = returns.rolling(window, min_periods=window).cov(benchmark_rets)
    var = benchmark_rets.rolling(window, min_periods=window).var()
    beta = cov / var
    return beta.replace([np.inf, -np.inf], np.nan)


PERCENT_METRICS = {
    "Annual return",
    "Cumulative returns",
    "Annual volatility",
    "Max drawdown",
    "Daily value at risk",
    "Alpha",
}

RATIO_METRICS = {
    "Sharpe ratio",
    "Calmar ratio",
    "Omega ratio",
    "Sortino ratio",
    "Tail ratio",
    "Stability",
    "Beta",
}


def format_stat(metric: str, value: float) -> str:
    if pd.isna(value):
        return ""
    if metric in PERCENT_METRICS:
        return f"{value:.2%}"
    if metric in RATIO_METRICS:
        return f"{value:.2f}"
    # Distribution moments etc.
    return f"{value:.2f}"


def month_returns(daily_rets: pd.Series) -> pd.Series:
    # Geometric monthly return
    return daily_rets.resample("ME").apply(lambda x: (1.0 + x).prod() - 1.0)


def drawdown_metadata(daily_rets: pd.Series) -> dict:
    dd = calc_drawdown(daily_rets)
    current_dd = float(dd.iloc[-1])
    max_dd = float(dd.min())
    trough_date = dd.idxmin()
    # Peak-to-trough start date (most recent peak prior to trough)
    wealth = calc_equity_curve(daily_rets, start=1.0)
    peak_to_date = wealth.loc[:trough_date].idxmax()
    return {
        "current_drawdown": current_dd,
        "max_drawdown": max_dd,
        "max_dd_peak_date": peak_to_date,
        "max_dd_trough_date": trough_date,
    }


def save_chart(path: Path, title: str, plot_fn):
    plt.figure()
    plot_fn()
    plt.title(title)
    plt.tight_layout()
    plt.savefig(path, dpi=200)
    plt.close()


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_stats_table(doc: Document, stats: pd.Series):
    # Convert to 2-col table: metric | value
    tbl = doc.add_table(rows=1, cols=2)
    hdr = tbl.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"
    for k, v in stats.items():
        row = tbl.add_row().cells
        row[0].text = str(k)
        row[1].text = format_stat(str(k), v)
    doc.add_paragraph()  # spacing


def add_bullet_list(doc: Document, items: list[str]):
    for item in items:
        if not item:
            continue
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph()


def build_executive_summary(
    returns: pd.Series,
    stats: pd.Series,
    benchmark_rets: pd.Series | None,
    rf_used: bool,
) -> list[str]:
    dd_meta = drawdown_metadata(returns)
    bullets: list[str] = []

    bullets.append(f"Cumulative return: {format_stat('Cumulative returns', stats.get('Cumulative returns', np.nan))} | Annualized return: {format_stat('Annual return', stats.get('Annual return', np.nan))}.")
    bullets.append(f"Risk: annualized vol {format_stat('Annual volatility', stats.get('Annual volatility', np.nan))}; max drawdown {format_stat('Max drawdown', stats.get('Max drawdown', np.nan))} (peak {dd_meta['max_dd_peak_date'].date()} → trough {dd_meta['max_dd_trough_date'].date()}).")

    sharpe = stats.get("Sharpe ratio", np.nan)
    sharpe_label = "excess" if rf_used else "rf=0"
    bullets.append(
        f"Sharpe ratio ({sharpe_label}): {format_stat('Sharpe ratio', sharpe)}; Sortino: {format_stat('Sortino ratio', stats.get('Sortino ratio', np.nan))}."
    )

    if benchmark_rets is not None:
        bullets.append(f"Benchmark-relative: beta {format_stat('Beta', stats.get('Beta', np.nan))}; alpha {format_stat('Alpha', stats.get('Alpha', np.nan))} (pyfolio estimate vs benchmark).")
    else:
        bullets.append("Benchmark-relative: not computed (no benchmark provided).")

    # Simple flags
    current_dd = dd_meta["current_drawdown"]
    if current_dd <= -0.10:
        bullets.append(f"Monitoring flag: currently in a drawdown of {current_dd:.2%}.")
    if stats.get("Max drawdown", 0.0) <= -0.20:
        bullets.append("Monitoring flag: max drawdown breached -20% threshold.")
    return bullets


def build_monitoring_notes(returns: pd.Series, benchmark_rets: pd.Series | None, rf_used: bool) -> list[str]:
    notes: list[str] = []
    dd_meta = drawdown_metadata(returns)

    # Worst days
    worst_days = returns.nsmallest(5)
    worst_days_txt = ", ".join([f"{d.date()}: {r:.2%}" for d, r in worst_days.items()])
    notes.append(f"Worst 5 days: {worst_days_txt}.")

    # Worst months
    mrets = month_returns(returns)
    if len(mrets) > 0:
        worst_months = mrets.nsmallest(3)
        worst_months_txt = ", ".join([f"{d.strftime('%Y-%m')}: {r:.2%}" for d, r in worst_months.items()])
        notes.append(f"Worst 3 months: {worst_months_txt}.")

    # Latest rolling risk metrics
    rvol = rolling_vol(returns, 63).dropna()
    rsh = rolling_sharpe(returns, 252).dropna()
    if len(rvol) > 0:
        notes.append(f"Latest 63d annualized vol: {rvol.iloc[-1]:.2%}.")
    if len(rsh) > 0:
        sharpe_label = "excess" if rf_used else "rf=0"
        notes.append(f"Latest 252d rolling Sharpe ({sharpe_label}): {rsh.iloc[-1]:.2f}.")

    notes.append(
        f"Current drawdown: {dd_meta['current_drawdown']:.2%}; max drawdown window: peak {dd_meta['max_dd_peak_date'].date()} → trough {dd_meta['max_dd_trough_date'].date()}."
    )

    if benchmark_rets is not None:
        rb = rolling_beta(returns, benchmark_rets, 252).dropna()
        if len(rb) > 0:
            notes.append(f"Latest 252d rolling beta vs benchmark: {rb.iloc[-1]:.2f}.")

    return notes


def add_bullets_section(doc: Document, title: str, bullets: list[str]):
    add_heading(doc, title, level=2)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph()


def add_image(doc: Document, img_path: Path, width_in: float = 6.5, caption: str | None = None):
    doc.add_picture(str(img_path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def compute_monitoring_flags(
    returns: pd.Series,
    benchmark_rets: pd.Series | None,
    dd_warn: float = -0.10,
    dd_breach: float = -0.20,
    rolling_sharpe_window: int = 252,
    rolling_sharpe_lookback: int = 126,
    sharpe_breach: float = 0.0,
    sharpe_breach_pct: float = 0.80,
    rolling_vol_window: int = 63,
    vol_breach: float = 0.20,
    rolling_beta_window: int = 252,
    beta_shift_breach: float = 0.30,
) -> list[str]:
    """Return human-readable monitoring flags suitable for an IC/monitoring pack.

    These are deliberately simple, transparent rules of thumb.
    """
    flags: list[str] = []
    dd = calc_drawdown(returns).dropna()
    if len(dd) == 0:
        return flags

    current_dd = float(dd.iloc[-1])
    max_dd = float(dd.min())

    # Drawdown flags
    if current_dd <= dd_breach:
        flags.append(f"Drawdown breach: current drawdown {current_dd:.2%} <= {dd_breach:.0%} threshold.")
    elif current_dd <= dd_warn:
        flags.append(f"Drawdown warning: current drawdown {current_dd:.2%} <= {dd_warn:.0%} threshold.")

    if max_dd <= dd_breach:
        flags.append(f"Max drawdown breach: max drawdown {max_dd:.2%} <= {dd_breach:.0%} threshold.")

    # Rolling Sharpe deterioration (rf=0 unless returns are excess)
    rsh = rolling_sharpe(returns, window=rolling_sharpe_window).dropna()
    if len(rsh) >= rolling_sharpe_lookback:
        tail = rsh.iloc[-rolling_sharpe_lookback:]
        pct_below = float((tail < sharpe_breach).mean())
        latest = float(tail.iloc[-1])
        if pct_below >= sharpe_breach_pct:
            flags.append(
                f"Sharpe deterioration: {pct_below:.0%} of last {rolling_sharpe_lookback} days had rolling Sharpe < {sharpe_breach:.2f} (latest {latest:.2f})."
            )
        elif latest < sharpe_breach:
            flags.append(
                f"Sharpe warning: latest rolling Sharpe {latest:.2f} < {sharpe_breach:.2f} (window {rolling_sharpe_window}d)."
            )

    # Vol spike
    rvol = rolling_vol(returns, window=rolling_vol_window).dropna()
    if len(rvol) > 0:
        latest_vol = float(rvol.iloc[-1])
        if latest_vol >= vol_breach:
            flags.append(
                f"Volatility breach: latest {rolling_vol_window}d annualized vol {latest_vol:.2%} >= {vol_breach:.0%} threshold."
            )

    # Beta regime shift
    if benchmark_rets is not None:
        rb = rolling_beta(returns, benchmark_rets, window=rolling_beta_window).dropna()
        if len(rb) > 0:
            latest_beta = float(rb.iloc[-1])
            # Baseline: median beta over history excluding the last window (to avoid leakage)
            baseline_series = rb.iloc[:-rolling_beta_window] if len(rb) > rolling_beta_window else rb
            baseline_beta = float(baseline_series.median())
            if abs(latest_beta - baseline_beta) >= beta_shift_breach:
                flags.append(
                    f"Beta regime shift: latest rolling beta {latest_beta:.2f} differs from baseline median {baseline_beta:.2f} by >= {beta_shift_breach:.2f}."
                )

    return flags


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--returns", default="data/manager_returns.csv")
    p.add_argument("--benchmark", default="data/benchmark_prices.csv")
    p.add_argument(
        "--risk_free",
        default="data/risk_free_daily.csv",
        help="CSV with daily risk-free returns (columns: date, rf). Used to compute excess returns for Sharpe/alpha.",
    )
    p.add_argument("--outdir", default="reports")
    p.add_argument("--report_name", default="manager_report.docx")
    p.add_argument("--title", default="Manager Performance & Risk Report")

    # Monitoring thresholds (rules of thumb; tune to your mandate)
    p.add_argument("--dd_warn", type=float, default=-0.10, help="Drawdown warning threshold (e.g., -0.10 for -10%)")
    p.add_argument("--dd_breach", type=float, default=-0.20, help="Drawdown breach threshold (e.g., -0.20 for -20%)")
    p.add_argument("--vol_breach", type=float, default=0.20, help="Annualized vol breach threshold (e.g., 0.20 for 20%)")
    p.add_argument("--sharpe_breach", type=float, default=0.0, help="Rolling Sharpe breach threshold")
    p.add_argument("--sharpe_lookback", type=int, default=126, help="Lookback days for rolling Sharpe deterioration test")
    p.add_argument("--sharpe_breach_pct", type=float, default=0.80, help="% of days below sharpe_breach to trigger deterioration flag")
    p.add_argument("--beta_shift_breach", type=float, default=0.30, help="Beta regime shift threshold vs baseline")
    args = p.parse_args()

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    charts_dir = outdir / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)

    # Load data
    returns_raw = load_returns(Path(args.returns))

    benchmark_rets_raw = None
    bench_path = Path(args.benchmark)
    if bench_path.exists():
        benchmark_rets_raw = load_benchmark_prices(bench_path)

    # Risk-free (optional but recommended)
    rf = None
    rf_path = Path(args.risk_free)
    if rf_path.exists():
        rf = load_risk_free_daily(rf_path)

    # Align series (raw) on the intersection of available dates
    common_idx = returns_raw.index
    if benchmark_rets_raw is not None:
        common_idx = common_idx.intersection(benchmark_rets_raw.index)
    if rf is not None:
        common_idx = common_idx.intersection(rf.index)

    returns_raw = returns_raw.reindex(common_idx)
    if benchmark_rets_raw is not None:
        benchmark_rets_raw = benchmark_rets_raw.reindex(common_idx)
    if rf is not None:
        rf = rf.reindex(common_idx)

    # Compute excess returns for risk-adjusted metrics
    returns = returns_raw - rf if rf is not None else returns_raw
    benchmark_rets = (
        (benchmark_rets_raw - rf) if (benchmark_rets_raw is not None and rf is not None) else benchmark_rets_raw
    )

    # Stats (pyfolio) — use *excess returns* if RF supplied
    if benchmark_rets is not None:
        stats = pf.timeseries.perf_stats(returns, factor_returns=benchmark_rets)
    else:
        stats = pf.timeseries.perf_stats(returns)

    stats.to_csv(outdir / "perf_stats.csv", header=False)

    # Also save a human-readable version
    stats_formatted = stats.copy()
    stats_formatted = stats_formatted.apply(lambda v: v)  # keep type
    (outdir / "perf_stats_formatted.csv").write_text(
        "metric,value\n" + "\n".join(
            f"{k},{format_stat(str(k), float(v) if pd.notna(v) else np.nan)}" for k, v in stats.items()
        )
        + "\n",
        encoding="utf-8",
    )

    # Build charts
    # Equity curve / drawdown should generally be shown on *raw* returns (what actually happened).
    mgr_equity = calc_equity_curve(returns_raw)
    mgr_dd = calc_drawdown(returns_raw)
    mgr_rvol = rolling_vol(returns, 63)
    mgr_rsharpe = rolling_sharpe(returns, 252)

    rb = None
    if benchmark_rets_raw is not None:
        bench_equity = calc_equity_curve(benchmark_rets_raw)
        rb = rolling_beta(returns, benchmark_rets, 252) if benchmark_rets is not None else None

    # 1) Equity curve
    equity_path = charts_dir / "equity_curve.png"
    def _plot_equity():
        mgr_equity.plot(label="Manager")
        if benchmark_rets_raw is not None:
            bench_equity.plot(label="Benchmark")
        plt.legend()
        plt.ylabel("Wealth (start=1.0)")
    save_chart(equity_path, "Equity Curve", _plot_equity)

    # 2) Drawdown
    dd_path = charts_dir / "drawdown.png"
    def _plot_dd():
        mgr_dd.plot()
        plt.ylabel("Drawdown")
    save_chart(dd_path, "Drawdown", _plot_dd)

    # 3) Rolling vol
    vol_path = charts_dir / "rolling_vol_63d.png"
    def _plot_vol():
        mgr_rvol.plot()
        plt.ylabel("Annualized Vol")
    save_chart(vol_path, "Rolling Volatility (63d)", _plot_vol)

    # 4) Rolling sharpe
    sharpe_path = charts_dir / "rolling_sharpe_252d.png"
    def _plot_sharpe():
        mgr_rsharpe.plot()
        plt.ylabel("Sharpe (annualized)")
    save_chart(sharpe_path, "Rolling Sharpe (252d)", _plot_sharpe)

    # 5) Rolling beta (if benchmark)
    beta_path = None
    if benchmark_rets is not None and rb is not None:
        beta_path = charts_dir / "rolling_beta_252d.png"
        def _plot_beta():
            rb.plot()
            plt.ylabel("Beta")
        save_chart(beta_path, "Rolling Beta vs Benchmark (252d)", _plot_beta)

    # 6) Return distribution
    hist_path = charts_dir / "return_hist.png"
    def _plot_hist():
        returns_raw.dropna().hist(bins=40)
        plt.xlabel("Return")
        plt.ylabel("Frequency")
    save_chart(hist_path, "Return Distribution", _plot_hist)

    # Create Word doc
    doc = Document()
    add_title(doc, args.title)
    doc.add_paragraph(f"Period: {returns_raw.index.min().date()} to {returns_raw.index.max().date()}")
    doc.add_paragraph(f"Observations: {len(returns_raw):,}")
    if rf is None:
        doc.add_paragraph("Risk-free rate: not supplied (Sharpe/alpha computed assuming rf=0).")
    else:
        doc.add_paragraph(f"Risk-free rate: daily series loaded from {rf_path.as_posix()} (Sharpe/alpha computed on excess returns).")
    doc.add_paragraph()

    # Executive summary + monitoring notes
    dd_meta = drawdown_metadata(returns_raw)
    worst_days = returns_raw.nsmallest(5)
    mrets = month_returns(returns_raw)
    worst_month = mrets.idxmin(), float(mrets.min())

    latest_vol = float(mgr_rvol.dropna().iloc[-1]) if len(mgr_rvol.dropna()) else np.nan
    latest_sharpe = float(mgr_rsharpe.dropna().iloc[-1]) if len(mgr_rsharpe.dropna()) else np.nan
    latest_beta = float(rb.dropna().iloc[-1]) if rb is not None and len(rb.dropna()) else np.nan

    bullets = [
        f"Cumulative return: {format_stat('Cumulative returns', stats.get('Cumulative returns', np.nan))} | Annualized return: {format_stat('Annual return', stats.get('Annual return', np.nan))}",
        f"Annualized vol: {format_stat('Annual volatility', stats.get('Annual volatility', np.nan))} | Sharpe ({'excess' if rf is not None else 'rf=0'}): {format_stat('Sharpe ratio', stats.get('Sharpe ratio', np.nan))}",
        f"Max drawdown: {format_stat('Max drawdown', stats.get('Max drawdown', np.nan))} (peak {dd_meta['max_dd_peak_date'].date()} → trough {dd_meta['max_dd_trough_date'].date()})",
    ]
    if rb is not None:
        bullets.append(
            f"Beta vs benchmark (latest 252d): {format_stat('Beta', latest_beta)} | Full-period beta: {format_stat('Beta', stats.get('Beta', np.nan))}"
        )
    bullets.append(
        f"Worst month: {worst_month[0].date()} ({worst_month[1]:.2%}) | Current drawdown: {dd_meta['current_drawdown']:.2%}"
    )

    add_heading(doc, "Executive Summary", level=1)
    add_bullets_section(doc, "Key takeaways", bullets)

    monitoring = [
        f"Current drawdown: {dd_meta['current_drawdown']:.2%}",
        f"Latest rolling vol (63d): {latest_vol:.2%} (annualized)" if pd.notna(latest_vol) else "Latest rolling vol (63d): n/a",
        f"Latest rolling Sharpe (252d, {'excess' if rf is not None else 'rf=0'}): {latest_sharpe:.2f}" if pd.notna(latest_sharpe) else "Latest rolling Sharpe (252d): n/a",
    ]
    if rb is not None:
        monitoring.append(
            f"Latest rolling beta (252d): {latest_beta:.2f}" if pd.notna(latest_beta) else "Latest rolling beta (252d): n/a"
        )
    monitoring.append("Worst 5 days (daily returns):")
    for dt, val in worst_days.items():
        monitoring.append(f"{dt.date()}: {val:.2%}")

    add_heading(doc, "Monitoring Notes", level=1)
    for line in monitoring:
        if line == "Worst 5 days (daily returns):":
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()

    # Monitoring flags (simple threshold-based alerts)
    flags = compute_monitoring_flags(
        returns,
        benchmark_rets,
        dd_warn=args.dd_warn,
        dd_breach=args.dd_breach,
        rolling_sharpe_window=252,
        rolling_sharpe_lookback=args.sharpe_lookback,
        sharpe_breach=args.sharpe_breach,
        sharpe_breach_pct=args.sharpe_breach_pct,
        rolling_vol_window=63,
        vol_breach=args.vol_breach,
        rolling_beta_window=252,
        beta_shift_breach=args.beta_shift_breach,
    )

    add_heading(doc, "Monitoring Flags", level=1)
    if len(flags) == 0:
        doc.add_paragraph("No threshold breaches detected based on the configured monitoring rules.")
    else:
        for f in flags:
            doc.add_paragraph(f, style="List Bullet")
    doc.add_paragraph()

    add_heading(doc, "Performance & Risk Summary", level=1)
    add_stats_table(doc, stats)

    add_heading(doc, "Charts", level=1)
    add_heading(doc, "Equity Curve", level=2)
    equity_caption = "Wealth index (start=1.0) aligned to common trading days."
    if benchmark_rets_raw is not None:
        equity_caption += " Benchmark shown as provided in benchmark_prices.csv."
    add_image(doc, equity_path, caption=equity_caption)

    add_heading(doc, "Drawdown", level=2)
    add_image(doc, dd_path)

    add_heading(doc, "Rolling Volatility", level=2)
    add_image(doc, vol_path)

    add_heading(doc, "Rolling Sharpe", level=2)
    add_image(doc, sharpe_path)

    if beta_path is not None:
        add_heading(doc, "Rolling Beta", level=2)
        add_image(doc, beta_path)

    add_heading(doc, "Return Distribution", level=2)
    add_image(doc, hist_path)

    report_path = outdir / args.report_name
    doc.save(report_path)

    print("Saved stats:", (outdir / "perf_stats.csv").resolve())
    print("Saved charts:", charts_dir.resolve())
    print("Saved Word report:", report_path.resolve())


if __name__ == "__main__":
    main()
